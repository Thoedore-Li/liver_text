import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_exam_doc(title, fill_blanks, multiple_choice, short_answer, answers, output_path):
    """
    创建考试文档
    """
    doc = Document()
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加填空题
    doc.add_paragraph("\n一、填空题（每空2分，共10分）")
    for question in fill_blanks:
        doc.add_paragraph(question)
    
    # 添加选择题
    doc.add_paragraph("\n二、单项选择题（每题2分，共10分）")
    for question in multiple_choice:
        doc.add_paragraph(question)
    
    # 添加简答题
    doc.add_paragraph("\n三、简答题（每题10分，共50分）")
    for question in short_answer:
        doc.add_paragraph(question)
    
    # 添加答案
    doc.add_page_break()  # 在新页面添加答案
    doc.add_paragraph(answers)
    
    # 保存文档
    doc.save(output_path)
    return output_path

# 修改保存路径，使用不同的文件名
specified_path = "D:\\download\\OneDrive\\Desktop\\text\\Liver_Exam_Paper2.docx"

# 试卷内容
exam_title = "肝病科住院医师考试试卷（二）"
fill_in_blanks_questions = [
    "1. 药物性肝损伤中，肝细胞毒性主要与______有关。",
    "2. NAFLD患者常伴发的代谢异常有______和______。",
    "3. 自身免疫性肝炎患者的主要影像学表现为______。",
    "4. 循环衰竭时，肝静脉压力升高可导致______。",
    "5. 药物性肝损伤的诊断依据包括______和______。"
]

multiple_choice_questions = [
    "1. 下列哪种药物最常引起胆汁淤积性肝损伤？\n   A. 甲氨蝶呤\n   B. 阿莫西林-克拉维酸\n   C. 阿司匹林\n   D. 阿托伐他汀",
    "2. NAFLD的主要病理表现是：\n   A. 小叶内纤维化\n   B. 胆管损伤\n   C. 肝细胞脂肪变性\n   D. 肝硬化",
    "3. 自身免疫性肝炎最常见的自身抗体是：\n   A. AMA\n   B. ANA\n   C. ANCA\n   D. ASMA",
    "4. 循环衰竭相关肝病中，以下哪种症状最为常见？\n   A. 食欲下降\n   B. 肝性脑病\n   C. 右上腹痛\n   D. 腹水",
    "5. 药物性肝损伤最重要的治疗原则是：\n   A. 停药\n   B. 手术\n   C. 激素治疗\n   D. 补充维生素"
]

short_answer_questions = [
    "1. 药物性肝损伤中慢性毒性的发生机制是什么？",
    "2. 描述NAFLD与心血管疾病之间的关系。",
    "3. 自身免疫性肝炎的诊断标准包括哪些？",
    "4. 循环衰竭相关肝病的预后评估有哪些指标？",
    "5. 药物性肝损伤患者的病因分析流程是怎样的？"
]

answers_content = """
答案：
一、填空题
1. 线粒体损伤
2. 高血压、糖尿病
3. 肝脏体积缩小
4. 肝淤血
5. 临床表现、实验室检查

二、单项选择题
1. B
2. C
3. B
4. C
5. A

三、简答题
1. 慢性毒性主要通过免疫介导和代谢产物蓄积，导致持续性肝细胞损伤和纤维化。
2. NAFLD患者心血管事件风险增加，与胰岛素抵抗、脂代谢紊乱和炎症因子释放有关。
3. 诊断标准包括：自身抗体阳性、血清IgG升高、组织学改变特征性、排除其他原因。
4. 预后评估指标：肝功能指标变化、凝血功能、器官功能状态、基础疾病控制情况。
5. 病因分析流程：详细用药史、排除其他原因、药物重新应用试验、因果关系评价。
"""

# 生成文档
result_path = create_exam_doc(
    exam_title,
    fill_in_blanks_questions,
    multiple_choice_questions,
    short_answer_questions,
    answers_content,
    specified_path
)

print(f"考试文档已生成，保存在：{result_path}") 