import sys
print(sys.executable)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_exam_doc(title, fill_blanks, multiple_choice, short_answer, answers, output_path):
    """
    创建考试文档
    """
    doc = Document()
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加填空题
    doc.add_paragraph("\n一、填空题")
    for question in fill_blanks:
        doc.add_paragraph(question)
    
    # 添加选择题
    doc.add_paragraph("\n二、单项选择题")
    for question in multiple_choice:
        doc.add_paragraph(question)
    
    # 添加简答题
    doc.add_paragraph("\n三、简答题")
    for question in short_answer:
        doc.add_paragraph(question)
    
    # 添加答案
    doc.add_page_break()  # 在新页面添加答案
    doc.add_paragraph(answers)
    
    # 保存文档
    doc.save(output_path)
    return output_path

# Adjusted save path to the specified location provided by the user
specified_path = "D:\\download\\OneDrive\\Desktop\\text\\Liver_Exam.docx"

# Define content for the exam based on the provided documents
exam_title = "肝病科住院医师考试试卷"
fill_in_blanks_questions = [
    "1. 药物性肝损伤的最常见病因药物是______。",
    "2. NAFLD的主要代谢机制包括______和脂质代谢紊乱。",
    "3. 自身免疫性肝炎的诊断标志物是______和______。",
    "4. 循环衰竭时肝窦的主要病理变化是______。",
    "5. 药物性肝损伤中慢性毒性的典型药物是______。"
]

multiple_choice_questions = [
    "1. 下列哪种药物最容易引起药物性肝损伤？\n   A. 阿莫西林-克拉维酸\n   B. 阿司匹林\n   C. 阿托伐他汀\n   D. 多奈哌齐",
    "2. NAFLD患者的病理特征不包括：\n   A. 肝细胞脂肪变性\n   B. 肝细胞再生障碍\n   C. 小叶性炎症\n   D. Mallory小体",
    "3. 自身免疫性肝炎的主要治疗是：\n   A. 抗生素\n   B. 免疫抑制剂\n   C. 抗病毒药物\n   D. 抗氧化剂",
    "4. 循环衰竭相关肝病的主要表现是：\n   A. 黄疸\n   B. 胰腺炎\n   C. 胆管扩张\n   D. 肝动脉栓塞",
    "5. 药物性肝损伤诊断的首要依据是：\n   A. 病理检查\n   B. 临床病史\n   C. 肝功能检查\n   D. 影像学检查"
]

short_answer_questions = [
    "1. 简述药物性肝损伤的主要类型及其临床特点。",
    "2. 描述NAFLD的关键发病机制及其相关危险因素。",
    "3. 解释自身免疫性肝炎与其他肝病的鉴别要点。",
    "4. 循环衰竭相关肝病的病理生理机制是什么？",
    "5. 药物性肝损伤的治疗原则包括哪些？"
]

answers_content = """
答案：
一、填空题
1. 阿莫西林-克拉维酸
2. 胰岛素抵抗
3. ANA和SMA
4. 肝窦毛细血管化
5. 甲氨蝶呤

二、单项选择题
1. A
2. B
3. B
4. A
5. B

三、简答题
1. 类型包括急性肝细胞损伤、胆汁淤积性损伤和混合型；特点分别为ALT升高、胆红素升高和两者均升高。
2. 发病机制包括胰岛素抵抗、脂质代谢紊乱；危险因素包括肥胖、糖尿病和代谢综合征。
3. 主要鉴别要点是自身抗体的检测和组织学特点。
4. 病理机制包括肝窦缺氧、炎症反应和纤维化。
5. 治疗原则包括停用可疑药物、支持治疗和应用特异性解毒剂（如N-乙酰半胱氨酸）。
"""

# Generate the document
result_path = create_exam_doc(
    exam_title,
    fill_in_blanks_questions,
    multiple_choice_questions,
    short_answer_questions,
    answers_content,
    specified_path
)

print(f"考试文档已生成，保存在：{result_path}")
