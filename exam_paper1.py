import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_exam_doc(title, fill_blanks, multiple_choice, short_answer, answers, output_path):
    """
    创建考试文档
    """
    doc = Document()
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加填空题
    doc.add_paragraph("\n一、填空题（每空2分，共10分）")
    for question in fill_blanks:
        doc.add_paragraph(question)
    
    # 添加选择题
    doc.add_paragraph("\n二、单项选择题（每题2分，共10分）")
    for question in multiple_choice:
        doc.add_paragraph(question)
    
    # 添加简答题
    doc.add_paragraph("\n三、简答题（每题10分，共50分）")
    for question in short_answer:
        doc.add_paragraph(question)
    
    # 添加答案
    doc.add_page_break()  # 在新页面添加答案
    doc.add_paragraph(answers)
    
    # 存文档
    doc.save(output_path)
    return output_path

# 修改保存路径，使用不同的文件名
specified_path = "D:\\download\\OneDrive\\Desktop\\text\\Liver_Exam_Paper1.docx"

# 试卷内容
exam_title = "肝病科住院医师考试试卷（一）"
fill_in_blanks_questions = [
    "1. 药物性肝损伤最常见的急性表现是______。",
    "2. NAFLD的组织学特点包括______和______。",
    "3. 自身免疫性肝炎常见的并发症包括______。",
    "4. 循环衰竭时肝脏发生的主要损伤机制是______。",
    "5. 药物性肝损伤的常见诱因包括______和______。"
]

multiple_choice_questions = [
    "1. 哪种检查对诊断药物性肝损伤最有帮助？\n   A. CT\n   B. MRI\n   C. 病理学活检\n   D. 临床病史",
    "2. NAFLD与下列哪项病理过程密切相关？\n   A. 胆汁淤积\n   B. 胰岛素抵抗\n   C. 酒精摄入\n   D. 胆道阻塞",
    "3. 自身免疫性肝炎患者中最常见的临床症状是：\n   A. 无症状\n   B. 黄疸\n   C. 腹痛\n   D. 发热",
    "4. 循环衰竭相关肝病的主要病理改变为：\n   A. 肝细胞增生\n   B. 肝窦毛细血管化\n   C. 胆管增生\n   D. 门静脉血栓形成",
    "5. 药物性肝损伤的早期治疗中最重要的是：\n   A. 病因治疗\n   B. 抗生素应用\n   C. 手术干预\n   D. 支持治疗"
]

short_answer_questions = [
    "1. 简述药物性肝损伤的高危人群及预防措施。",
    "2. 描述NAFLD中肝细胞脂肪变性的形成机制。",
    "3. 自身免疫性肝炎的治疗目标是什么？常用药物有哪些？",
    "4. 循环衰竭相关肝病的主要并发症及其临床意义。",
    "5. 如何评估药物性肝损伤的严重程度？"
]

answers_content = """
答案：
一、填空题
1. 转氨酶升高
2. 肝细胞脂肪变性、炎症浸润
3. 肝硬化
4. 缺血缺氧
5. 药物过量、特异体质

二、单项选择题
1. D
2. B
3. B
4. B
5. D

三、简答题
1. 高危人群：老年人、肝病患者、多药联用者；预防措施：避免不必要用药、定期监测肝功能、注意药物相互作用。
2. 胰岛素抵抗导致脂肪酸代谢异常，游离脂肪酸在肝细胞内堆积，同时氧化应激和线粒体功能障碍加重脂肪变性。
3. 治疗目标：缓解症状、降低转氨酶、改善预后；常用药物：泼尼松、硫唑嘌呤等免疫抑制剂。
4. 主要并发症：肝功能衰竭、凝血功能障碍、肾功能不全；临床意义：影响预后、增加病死率。
5. 评估指标：肝功能指标、凝血功能、临床症状体征、影像学改变等综合评估。
"""

# 生成文档
result_path = create_exam_doc(
    exam_title,
    fill_in_blanks_questions,
    multiple_choice_questions,
    short_answer_questions,
    answers_content,
    specified_path
)

print(f"考试文档已生成，保存在：{result_path}") 