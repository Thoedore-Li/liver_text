import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_exam_doc(title, fill_blanks, multiple_choice, short_answer, answers, output_path):
    """
    创建考试文档
    """
    doc = Document()
    
    # 添加标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加填空题
    doc.add_paragraph("\n一、填空题（每空2分，共10分）")
    for question in fill_blanks:
        doc.add_paragraph(question)
    
    # 添加选择题
    doc.add_paragraph("\n二、单项选择题（每题2分，共10分）")
    for question in multiple_choice:
        doc.add_paragraph(question)
    
    # 添加简答题
    doc.add_paragraph("\n三、简答题（每题10分，共50分）")
    for question in short_answer:
        doc.add_paragraph(question)
    
    # 添加答案
    doc.add_page_break()  # 在新页面添加答案
    doc.add_paragraph(answers)
    
    # 保存文档
    doc.save(output_path)
    return output_path

# 修改保存路径，使用不同的文件名
specified_path = "D:\\download\\OneDrive\\Desktop\\text\\Liver_Exam_Paper3.docx"

# 试卷内容
exam_title = "肝病科住院医师考试试卷（三）"
fill_in_blanks_questions = [
    "1. 药物性肝损伤中，肝小叶中心区坏死多与______相关。",
    "2. NAFLD的组织学进展包括______和______。",
    "3. 自身免疫性肝炎最常见的病理学表现为______。",
    "4. 循环衰竭相关肝病中，肝窦的主要病理变化是______。",
    "5. 药物性肝损伤中，胆汁淤积型的特征是______。"
]

multiple_choice_questions = [
    "1. 哪种药物的代谢与肝细胞线粒体毒性密切相关？\n   A. 四环素\n   B. 对乙酰氨基酚\n   C. 阿莫西林\n   D. 头孢菌素",
    "2. NAFLD的危险因素不包括：\n   A. 肥胖\n   B. 胰岛素抵抗\n   C. 慢性病毒性肝炎\n   D. 代谢综合征",
    "3. 自身免疫性肝炎患者常见的血清学异常是：\n   A. ALP升高\n   B. ALT升高\n   C. 淋巴细胞减少\n   D. 中性粒细胞减少",
    "4. 循环衰竭相关肝病的主要病理表现为：\n   A. 胆汁淤积\n   B. 肝窦扩张\n   C. 肝脏结节形成\n   D. 门静脉闭塞",
    "5. 药物性肝损伤的常见诱因是：\n   A. 饮酒\n   B. 免疫抑制剂\n   C. 抗生素\n   D. 病毒感染"
]

short_answer_questions = [
    "1. 简述药物性肝损伤的早期症状和体征。",
    "2. NAFLD患者如何通过生活方式干预减轻疾病进展？",
    "3. 自身免疫性肝炎与胆汁淤积性肝病的主要鉴别点是什么？",
    "4. 循环衰竭相关肝病的治疗策略有哪些？",
    "5. 药物性肝损伤中肝功能恢复的主要监测指标有哪些？"
]

answers_content = """
答案：
一、填空题
1. 药物代谢
2. 炎症、纤维化
3. 界面性肝炎
4. 肝窦扩张
5. 血清ALP和γ-GT升高

二、单项选择题
1. B
2. C
3. B
4. B
5. C

三、简答题
1. 早期症状：乏力、食欲下降、恶心；体征：黄疸、肝大、压痛。
2. 生活方式干预：控制饮食、规律运动、戒酒、控制体重、血糖管理。
3. 鉴别点：血清学指标、组织学特征、治疗反应、疾病进展特点。
4. 治疗策略：改善循环、保护器官功能、预防并发症、原发病治疗。
5. 监测指标：转氨酶、胆红素、凝血功能、临床症状改善情况。
"""

# 生成文档
result_path = create_exam_doc(
    exam_title,
    fill_in_blanks_questions,
    multiple_choice_questions,
    short_answer_questions,
    answers_content,
    specified_path
)

print(f"考试文档已生成，保存在：{result_path}") 